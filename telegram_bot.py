from telegram import Update, Document
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes
from api import telegram_api
from query_rag import answer_question
from data_rag import add_document_to_vector_store

import os
from pathlib import Path
import glob
import re

# Удалены все переменные и функции, связанные с Google API

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message or not update.message.text:
        return
    text = update.message.text
    response = answer_question(text)
    await update.message.reply_text(response)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.message:
        await update.message.reply_text("Привет! Задай вопрос или пришли документ для добавления в базу.")

async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message or not update.message.document:
        return
    document = update.message.document
    file_name = document.file_name
    if not file_name:
        if update.message:
            await update.message.reply_text("Файл должен иметь имя.")
        return

    # Скачиваем файл
    file_path = Path("downloads") / file_name
    file_path.parent.mkdir(exist_ok=True)

    file = await context.bot.get_file(document.file_id)
    await file.download_to_drive(str(file_path))

    n_chunks = 0
    # Чтение файла и индексация по частям
    if file_name.endswith(".txt"):
        with file_path.open(encoding="utf-8") as f:
            for i, line in enumerate(f):
                line = line.strip()
                if line:
                    n_chunks += add_document_to_vector_store(line, f"{file_name} (строка {i+1})")
    elif file_name.endswith(".pdf"):
        from PyPDF2 import PdfReader
        reader = PdfReader(str(file_path))
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                n_chunks += add_document_to_vector_store(page_text, f"{file_name} (стр. {i+1})")
    elif file_name.endswith(".docx"):
        from docx import Document as DocxDocument
        doc = DocxDocument(str(file_path))
        for i, p in enumerate(doc.paragraphs):
            para_text = p.text.strip()
            if para_text:
                n_chunks += add_document_to_vector_store(para_text, f"{file_name} (абзац {i+1})")
    else:
        if update.message:
            await update.message.reply_text("Поддерживаются только .txt, .pdf, .docx")
        return

    if update.message:
        await update.message.reply_text(f"Документ добавлен. Создано {n_chunks} фрагментов.")

def index_all_documents():
    supported_exts = [".txt", ".pdf", ".docx"]
    downloads_path = Path("downloads")
    if not downloads_path.exists():
        return
    for file_path in downloads_path.iterdir():
        if file_path.suffix.lower() not in supported_exts:
            continue
        try:
            if file_path.suffix.lower() == ".txt":
                content = file_path.read_text(encoding="utf-8")
            elif file_path.suffix.lower() == ".pdf":
                from PyPDF2 import PdfReader
                reader = PdfReader(str(file_path))
                content = "\n".join(page.extract_text() for page in reader.pages)
            elif file_path.suffix.lower() == ".docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(str(file_path))
                content = "\n".join(p.text for p in doc.paragraphs)
            else:
                continue
            add_document_to_vector_store(content, file_path.name)
        except Exception as e:
            print(f"Ошибка при индексации {file_path.name}: {e}")

def main():
    index_all_documents()
    app = ApplicationBuilder().token(telegram_api).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
    app.run_polling()

if __name__ == "__main__":
    main()
